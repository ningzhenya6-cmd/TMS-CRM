"""
文档导出 — Word (.docx) + PDF 文档生成
支持两种报告类型：risk（学业风险分析）, preparation（行前准备规划）
"""
import json
import logging
import os
from io import BytesIO

logger = logging.getLogger(__name__)

# ── Word ──
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# ── PDF ──
from fpdf import FPDF


# ══════════════════════════════════════════════════════════════
#  共用工具
# ══════════════════════════════════════════════════════════════

_RISK_COLORS = {
    "low": ("#16a34a", "低"),
    "medium": ("#ca8a04", "中"),
    "high": ("#dc2626", "高"),
}

_PRIORITY_LABELS = {
    "high": "🔴 高优先级",
    "medium": "🟡 中优先级",
    "low": "🟢 低优先级",
}

_READINESS_LABELS = {
    "low": "低",
    "medium": "中",
    "high": "高",
}


def _get_lead_info(lead):
    """提取学生信息，兼容 None"""
    if not lead:
        return {"name": "未知", "phone": "", "country": "", "grade": ""}
    return {
        "name": lead.get("name", "未知"),
        "phone": lead.get("phone", ""),
        "country": lead.get("country", ""),
        "grade": lead.get("grade", ""),
    }


def _parse_report_data(report_json):
    """解析 report_json，返回 dict 或 None"""
    if not report_json:
        return None
    try:
        return json.loads(report_json) if isinstance(report_json, str) else report_json
    except (json.JSONDecodeError, TypeError):
        return None


# ══════════════════════════════════════════════════════════════
#  Word 文档
# ══════════════════════════════════════════════════════════════


def _set_cell_shading(cell, color):
    """设置表格单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color, qn("w:val"): "clear"},
    )
    shading.append(shading_elem)


def _add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        _set_cell_shading(cell, "1e40af")
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # 数据行
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        # 交替行背景
        if ri % 2 == 1:
            for ci in range(len(headers)):
                _set_cell_shading(table.rows[ri + 1].cells[ci], "f1f5f9")


def _build_risk_docx(report, report_data, lead):
    """生成学业风险分析 Word 文档"""
    doc = Document()
    li = _get_lead_info(lead)

    # ── 标题 ──
    title = doc.add_heading("学业风险分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 学生信息 ──
    doc.add_heading("学生信息", level=1)
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("学生姓名", li["name"]),
        ("联系方式", li["phone"]),
        ("目标院校", f"{report.get('target_school', '')} - {report.get('target_major', '')}"),
        ("意向国家", report.get("target_country", "")),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for run in info_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()  # spacing

    if not report_data:
        doc.add_paragraph("报告数据为空")
        return _docx_bytes(doc)

    # ── 整体风险等级 ──
    risk_level = report_data.get("overall_risk", "medium")
    risk_info = _RISK_COLORS.get(risk_level, ("#6b7280", "未知"))
    doc.add_heading("整体评估", level=1)
    p = doc.add_paragraph()
    run = p.add_run(f"风险等级：{risk_info[1]}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*[int(risk_info[0][i:i+2], 16) for i in (1, 3, 5)])

    # ── 学生概况 ──
    overview = report_data.get("profile_overview", "")
    if overview:
        doc.add_heading("学生概况", level=2)
        doc.add_paragraph(overview)

    # ── 目标院校要求 ──
    req = report_data.get("program_requirements", {})
    if req:
        doc.add_heading("目标院校要求", level=2)
        for key, label in [
            ("gpa_requirement", "GPA 要求"),
            ("language_requirement", "语言要求"),
            ("assessment_format", "考核形式"),
            ("curriculum_highlights", "课程特色"),
        ]:
            val = req.get(key)
            if val:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}：")
                run.bold = True
                p.add_run(str(val))

        prereqs = req.get("prerequisites", [])
        if prereqs:
            doc.add_paragraph("前置修读科目：")
            for pr in prereqs:
                doc.add_paragraph(pr, style="List Bullet")

    # ── 差距分析表 ──
    gaps = report_data.get("gap_analysis", [])
    if gaps:
        doc.add_heading("差距分析", level=2)
        gap_rows = []
        for g in gaps:
            risk_lvl = g.get("risk", "medium")
            risk_label = _RISK_COLORS.get(risk_lvl, ("#6b7280", "未知"))[1]
            gap_rows.append([
                g.get("dimension", ""),
                g.get("current", ""),
                g.get("required", ""),
                g.get("gap", ""),
                risk_label,
            ])
        _add_styled_table(doc, ["维度", "当前", "要求", "差距", "风险"], gap_rows)
        doc.add_paragraph()

    # ── 建议方案 ──
    recs = report_data.get("recommendations", [])
    if recs:
        doc.add_heading("建议方案", level=2)
        for r in recs:
            priority = r.get("priority", "medium")
            label = _PRIORITY_LABELS.get(priority, f"[{priority}]")
            p = doc.add_paragraph()
            run = p.add_run(f"{label}：{r.get('action', '')}")
            run.bold = True
            timeline = r.get("timeline", "")
            impact = r.get("expected_impact", "")
            if timeline:
                doc.add_paragraph(f"  时间线：{timeline}")
            if impact:
                doc.add_paragraph(f"  预期效果：{impact}")

    # ── 顾问沟通建议 ──
    tips = report_data.get("consultant_tips", "")
    if tips:
        doc.add_heading("顾问沟通建议", level=2)
        doc.add_paragraph(tips)

    return _docx_bytes(doc)


def _build_preparation_docx(report, report_data, lead):
    """生成行前准备规划 Word 文档"""
    doc = Document()
    li = _get_lead_info(lead)

    # ── 标题 ──
    title = doc.add_heading("行前学业准备规划报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 学生信息 ──
    doc.add_heading("学生信息", level=1)
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("学生姓名", li["name"]),
        ("联系方式", li["phone"]),
        ("目标院校", report.get("target_school", "")),
        ("目标专业", report.get("target_major", "")),
        ("意向国家", report.get("target_country", "")),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for run in info_table.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True

    doc.add_paragraph()

    if not report_data:
        doc.add_paragraph("报告数据为空")
        return _docx_bytes(doc)

    # ── 项目总览 ──
    overview = report_data.get("program_overview", "")
    if overview:
        doc.add_heading("项目总览", level=1)
        doc.add_paragraph(overview)

    # ── 课程级准备规划 ──
    courses = report_data.get("courses", [])
    if courses:
        doc.add_heading("课程准备规划", level=1)
        for c in courses:
            # 课程名 heading
            code = c.get("course_code", "")
            name = c.get("course_name", "")
            heading_text = f"{code} {name}".strip()
            doc.add_heading(heading_text, level=2)

            # 课程描述
            desc = c.get("course_description", "")
            if desc:
                p = doc.add_paragraph()
                run = p.add_run("课程描述：")
                run.bold = True
                p.add_run(desc)

            # 核心主题
            topics = c.get("core_topics", [])
            if topics:
                doc.add_paragraph("核心主题：" + "、".join(topics))

            # 前置要求
            prereq = c.get("prerequisites_expected", "")
            if prereq:
                p = doc.add_paragraph()
                run = p.add_run("前置要求：")
                run.bold = True
                p.add_run(prereq)

            # Readiness
            readiness = c.get("student_readiness", "medium")
            readiness_label = _READINESS_LABELS.get(readiness, readiness)
            p = doc.add_paragraph()
            run = p.add_run(f"学生准备度：{readiness_label}")
            run.bold = True

            analysis = c.get("readiness_analysis", "")
            if analysis:
                doc.add_paragraph(f"分析：{analysis}")

            # 准备行动
            actions = c.get("preparation_actions", [])
            if actions:
                doc.add_paragraph("准备行动建议：")
                for a in actions:
                    doc.add_paragraph(a, style="List Bullet")

            # 考核形式
            assess = c.get("assessment_format", "")
            if assess:
                p = doc.add_paragraph()
                run = p.add_run("考核形式：")
                run.bold = True
                p.add_run(assess)

            # 推荐资源
            resources = c.get("recommended_resources", [])
            if resources:
                doc.add_paragraph("推荐学习资源：")
                for r in resources:
                    doc.add_paragraph(r, style="List Bullet")

    # ── 整体时间线 ──
    timeline = report_data.get("overall_timeline", "")
    if timeline:
        doc.add_heading("整体准备时间线", level=1)
        doc.add_paragraph(timeline)

    # ── 优先关注点 ──
    focus = report_data.get("priority_focus", "")
    if focus:
        doc.add_heading("优先关注点", level=1)
        doc.add_paragraph(focus)

    # ── 顾问建议 ──
    notes = report_data.get("advisor_notes", "")
    if notes:
        doc.add_heading("顾问建议", level=1)
        doc.add_paragraph(notes)

    return _docx_bytes(doc)


def _docx_bytes(doc) -> bytes:
    """Document → bytes"""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_unified_docx(report, report_data, lead):
    """生成统一学业风险与规划报告 Word 文档"""
    doc = Document()
    li = _get_lead_info(lead)

    title = doc.add_heading("学业风险与规划报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告日期：{report.get('created_at', '')[:10] or '未知'}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not report_data:
        doc.add_paragraph("报告数据为空")
        return _docx_bytes(doc)

    # 学生基础画像
    sp = report_data.get("student_profile", {})
    if sp:
        doc.add_heading("学生基础画像", level=1)
        doc.add_paragraph(sp.get("background", ""))
        p = doc.add_paragraph()
        run = p.add_run("目标：")
        run.bold = True
        p.add_run(sp.get("target", ""))
        ki = sp.get("key_info", {})
        if ki:
            doc.add_paragraph()
            for k, v in ki.items():
                doc.add_paragraph(f"{k}：{v}", style="List Bullet")

    # 目标院校概览
    po = report_data.get("program_overview", "")
    if po:
        doc.add_heading("目标院校专业概览", level=1)
        doc.add_paragraph(po)

    # 逐课程分析
    courses = report_data.get("course_analysis", [])
    if courses:
        doc.add_heading("核心课程逐门分析", level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        for i, label in enumerate(["课程", "学分", "考核形式", "前置要求", "学生掌握", "差距分析", "建议补充"]):
            hdr[i].text = label
        for c in courses:
            row = table.add_row().cells
            row[0].text = c.get("course_name", "")
            row[1].text = str(c.get("credits", ""))
            row[2].text = c.get("assessment", "")
            row[3].text = c.get("prerequisites", "")
            row[4].text = c.get("student_readiness", "")
            row[5].text = c.get("gap", "")
            row[6].text = c.get("supplement_action", "")
        doc.add_paragraph()

    # 差距分析
    gaps = report_data.get("gap_analysis", [])
    if gaps:
        doc.add_heading("多维差距分析", level=1)
        for g in gaps:
            p = doc.add_paragraph()
            run = p.add_run(f"{g.get('dimension', '')}（{g.get('risk', 'medium')}）")
            run.bold = True
            doc.add_paragraph(f"当前：{g.get('current', '')}")
            doc.add_paragraph(f"要求：{g.get('required', '')}")
            doc.add_paragraph(f"差距：{g.get('gap', '')}")
            if g.get("improvement_strategy"):
                doc.add_paragraph(f"改进策略：{g['improvement_strategy']}")
            doc.add_paragraph()

    # 准备计划
    plans = report_data.get("preparation_plan", [])
    if plans:
        doc.add_heading("准备计划", level=1)
        for plan in plans:
            p = doc.add_paragraph()
            run = p.add_run(f"{plan.get('phase', '')}")
            run.bold = True
            if plan.get("timeline"):
                p.add_run(f" · {plan['timeline']}")
            if plan.get("goal"):
                doc.add_paragraph(f"目标：{plan['goal']}")
            for task in plan.get("tasks", []):
                doc.add_paragraph(task, style="List Bullet")

    # 综合评估
    oa = report_data.get("overall_assessment", "")
    if oa:
        doc.add_heading("综合评估", level=1)
        doc.add_paragraph(oa)
        rl = report_data.get("risk_level", "medium")
        p = doc.add_paragraph()
        run = p.add_run(f"整体风险：{rl}")
        run.bold = True

    # 建议
    recs = report_data.get("recommendations", [])
    if recs:
        doc.add_heading("具体建议", level=1)
        for i, rec in enumerate(recs, 1):
            text = rec if isinstance(rec, str) else rec.get("action", str(rec))
            doc.add_paragraph(f"{i}. {text}")

    return _docx_bytes(doc)



def _build_growth_docx(report, report_data, lead):
    """生成整体学情报告 Word 文档"""
    from docx import Document
    doc = Document()
    li = {"name": lead.get("name","未知")} if lead else {"name":"未知"}
    
    doc.add_heading(report_data.get("report_title", "阶段性学情综合报告"), level=0)
    
    si = report_data.get("student_info", {})
    if si:
        doc.add_heading("学生信息", level=1)
        for k, v in si.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{k}：")
            run.bold = True
            p.add_run(str(v))
    
    bl = report_data.get("baseline", "")
    if bl:
        doc.add_heading("入学前基础概况", level=1)
        doc.add_paragraph(bl)
    
    ls = report_data.get("learning_summary", "")
    if ls:
        doc.add_heading("学习历程总结", level=1)
        doc.add_paragraph(ls)
    
    for title, field, color in [("进步方面", "progress", "green"), ("待改进方面", "weaknesses", "red")]:
        items = report_data.get(field, [])
        if items:
            doc.add_heading(title, level=1)
            for item in items:
                doc.add_paragraph(f"{item.get('aspect','')}：{item.get('detail','')}")
    
    et = report_data.get("exam_trend", "")
    if et:
        doc.add_heading("考试成绩趋势", level=1)
        doc.add_paragraph(et)
    
    dims = report_data.get("dimensions", {})
    if dims:
        doc.add_heading("多维度评估", level=1)
        labels = {"academic":"学术","language":"语言","psychology":"心理","adaptation":"适应","planning":"规划"}
        for key, dim in dims.items():
            p = doc.add_paragraph()
            run = p.add_run(f"{labels.get(key,key)}（{dim.get('score','')}）：")
            run.bold = True
            p.add_run(dim.get("note",""))
    
    oa = report_data.get("overall_assessment", "")
    if oa:
        doc.add_heading("综合评估", level=1)
        doc.add_paragraph(oa)
    
    recs = report_data.get("recommendations", [])
    if recs:
        doc.add_heading("后续建议", level=1)
        for i, rec in enumerate(recs, 1):
            doc.add_paragraph(f"{i}. {rec}")
    
    from io import BytesIO
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

def generate_docx(report, lead, report_type="risk") -> bytes:
    """
    生成 Word 文档

    Args:
        report: consulting_reports 行 dict（含 report_json）
        lead: leads 行 dict
        report_type: 'risk' | 'preparation' | 'unified'

    Returns:
        .docx 文件 bytes
    """
    report_data = _parse_report_data(report.get("report_json", ""))

    if report_type == "preparation":
        return _build_preparation_docx(report, report_data, lead)
    if report_type == "growth_overall":
        return _build_growth_docx(report, report_data, lead)
    if report_type == "unified":
        return _build_unified_docx(report, report_data, lead)
    return _build_risk_docx(report, report_data, lead)


# ══════════════════════════════════════════════════════════════
#  PDF 文档
# ══════════════════════════════════════════════════════════════


def _find_chinese_font():
    """自动检测可用中文字体路径，支持 macOS / Linux"""
    candidates = [
        # macOS
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        # Linux — WenQuanYi
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        # Linux — Noto
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        # Linux — 其他常见路径
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
    ]
    for p in candidates:
        if os.path.exists(p):
            return p
    return None


class _ReportPDF(FPDF):
    """自定义 PDF，自动检测中文字体，页脚加页码"""

    def __init__(self):
        super().__init__()
        self._chinese_ok = False
        self._cn_font = "Helvetica"
        font_path = _find_chinese_font()
        if font_path:
            try:
                self.add_font("CJK", "", font_path)
                # 粗体复用同一字体（fpdf2 支持）
                self.add_font("CJK", "B", font_path)
                self._chinese_ok = True
                self._cn_font = "CJK"
            except Exception as e:
                logger.error('Failed to add Chinese font to PDF, falling back to Helvetica', extra={'error': str(e), 'font_path': font_path})

    def footer(self):
        self.set_y(-15)
        if self._chinese_ok:
            self.set_font(self._cn_font, "", 8)
        else:
            self.set_font("Helvetica", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def _pdf_set_cn(pdf):
    """切换到中文字体"""
    if pdf._chinese_ok:
        return pdf._cn_font
    return "Helvetica"


def _pdf_section_header(pdf, title):
    """PDF 章节标题"""
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 14)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
    pdf.set_draw_color(30, 64, 175)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
    pdf.ln(4)


def _pdf_sub_header(pdf, title):
    """PDF 子章节标题"""
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 11)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 8, title, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)


def _pdf_body(pdf, text):
    """PDF 正文"""
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "", 10)
    pdf.set_text_color(50, 50, 50)
    pdf.multi_cell(0, 6, text)
    pdf.ln(2)


def _pdf_bold_line(pdf, label, value=""):
    """标签: 值"""
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 10)
    pdf.set_text_color(50, 50, 50)
    w = pdf.get_string_width(label) + 2
    pdf.cell(w, 7, label)
    if value:
        pdf.set_font(font, "", 10)
        pdf.cell(0, 7, value, new_x="LMARGIN", new_y="NEXT")
    else:
        pdf.ln(7)


def _pdf_bullet(pdf, text, indent=10):
    """带缩进的列表项"""
    font = _pdf_set_cn(pdf)
    pdf.set_x(pdf.get_x() + indent)
    pdf.set_font(font, "", 10)
    pdf.set_text_color(50, 50, 50)
    pdf.cell(5, 6, chr(8226))  # bullet
    pdf.multi_cell(0, 6, text)
    pdf.ln(1)


def _build_risk_pdf(report, report_data, lead):
    """生成学业风险分析 PDF"""
    pdf = _ReportPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    li = _get_lead_info(lead)

    # 标题
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 22)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(0, 15, "Academic Risk Analysis", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font, "", 16)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "学业风险分析报告", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # 学生信息
    _pdf_section_header(pdf, "Student Information")
    info_items = [
        ("Name:", li["name"]),
        ("Phone:", li["phone"]),
        ("Target:", f"{report.get('target_school', '')} - {report.get('target_major', '')}"),
        ("Country:", report.get("target_country", "")),
    ]
    for label, value in info_items:
        _pdf_bold_line(pdf, label, value)

    if not report_data:
        _pdf_body(pdf, "Report data is empty")
        return _pdf_bytes(pdf)

    # Overall risk
    risk_level = report_data.get("overall_risk", "medium")
    risk_info = _RISK_COLORS.get(risk_level, ("#6b7280", "Unknown"))
    _pdf_section_header(pdf, "Overall Assessment")
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 14)
    pdf.set_text_color(*[int(risk_info[0][i:i+2], 16) for i in (1, 3, 5)])
    pdf.cell(0, 10, f"Risk Level: {risk_info[1]}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Profile overview
    overview = report_data.get("profile_overview", "")
    if overview:
        _pdf_section_header(pdf, "Profile Overview")
        _pdf_body(pdf, overview)

    # Program requirements
    req = report_data.get("program_requirements", {})
    if req:
        _pdf_section_header(pdf, "Program Requirements")
        for key, label in [
            ("gpa_requirement", "GPA Requirement:"),
            ("language_requirement", "Language Requirement:"),
            ("assessment_format", "Assessment Format:"),
            ("curriculum_highlights", "Curriculum Highlights:"),
        ]:
            val = req.get(key)
            if val:
                _pdf_bold_line(pdf, label, str(val))

        prereqs = req.get("prerequisites", [])
        if prereqs:
            _pdf_bold_line(pdf, "Prerequisites:")
            for pr in prereqs:
                _pdf_bullet(pdf, pr)

    # Gap analysis
    gaps = report_data.get("gap_analysis", [])
    if gaps:
        _pdf_section_header(pdf, "Gap Analysis")
        for g in gaps:
            risk_lvl = g.get("risk", "medium")
            risk_label = _RISK_COLORS.get(risk_lvl, ("", "?"))[1]
            _pdf_sub_header(pdf, f"{g.get('dimension', '')} [{risk_label}]")
            _pdf_bold_line(pdf, "  Current:", str(g.get("current", "")))
            _pdf_bold_line(pdf, "  Required:", str(g.get("required", "")))
            _pdf_bold_line(pdf, "  Gap:", str(g.get("gap", "")))
            detail = g.get("detail", "")
            if detail:
                _pdf_body(pdf, f"  {detail}")

    # Recommendations
    recs = report_data.get("recommendations", [])
    if recs:
        _pdf_section_header(pdf, "Recommendations")
        for r in recs:
            priority = r.get("priority", "medium")
            label = _PRIORITY_LABELS.get(priority, f"[{priority}]")
            _pdf_sub_header(pdf, label)
            _pdf_bold_line(pdf, "  Action:", str(r.get("action", "")))
            timeline = r.get("timeline", "")
            if timeline:
                _pdf_bold_line(pdf, "  Timeline:", timeline)
            impact = r.get("expected_impact", "")
            if impact:
                _pdf_bold_line(pdf, "  Expected Impact:", impact)

    # Consultant tips
    tips = report_data.get("consultant_tips", "")
    if tips:
        _pdf_section_header(pdf, "Consultant Tips")
        _pdf_body(pdf, tips)

    return _pdf_bytes(pdf)


def _build_preparation_pdf(report, report_data, lead):
    """生成行前准备规划 PDF"""
    pdf = _ReportPDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    li = _get_lead_info(lead)

    # 标题
    font = _pdf_set_cn(pdf)
    pdf.set_font(font, "B", 22)
    pdf.set_text_color(30, 64, 175)
    pdf.cell(0, 15, "Pre-Departure Preparation Plan", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font, "", 16)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, "行前学业准备规划报告", align="C", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # 学生信息
    _pdf_section_header(pdf, "Student Information")
    info_items = [
        ("Name:", li["name"]),
        ("Phone:", li["phone"]),
        ("Target School:", report.get("target_school", "")),
        ("Target Major:", report.get("target_major", "")),
        ("Country:", report.get("target_country", "")),
    ]
    for label, value in info_items:
        _pdf_bold_line(pdf, label, value)

    if not report_data:
        _pdf_body(pdf, "Report data is empty")
        return _pdf_bytes(pdf)

    # 项目总览
    overview = report_data.get("program_overview", "")
    if overview:
        _pdf_section_header(pdf, "Program Overview")
        _pdf_body(pdf, overview)

    # 课程级准备规划
    courses = report_data.get("courses", [])
    if courses:
        _pdf_section_header(pdf, "Course Preparation Plan")
        for c in courses:
            code = c.get("course_code", "")
            name = c.get("course_name", "")
            heading = f"{code} {name}".strip()
            _pdf_sub_header(pdf, heading)

            desc = c.get("course_description", "")
            if desc:
                _pdf_bold_line(pdf, "Description:", desc)

            topics = c.get("core_topics", [])
            if topics:
                _pdf_bold_line(pdf, "Core Topics:", ", ".join(topics))

            prereq = c.get("prerequisites_expected", "")
            if prereq:
                _pdf_bold_line(pdf, "Prerequisites:", prereq)

            readiness = c.get("student_readiness", "medium")
            readiness_label = _READINESS_LABELS.get(readiness, readiness)
            _pdf_bold_line(pdf, "Student Readiness:", readiness_label)

            analysis = c.get("readiness_analysis", "")
            if analysis:
                _pdf_body(pdf, analysis)

            actions = c.get("preparation_actions", [])
            if actions:
                _pdf_bold_line(pdf, "Preparation Actions:")
                for a in actions:
                    _pdf_bullet(pdf, a)

            assess = c.get("assessment_format", "")
            if assess:
                _pdf_bold_line(pdf, "Assessment:", assess)

            resources = c.get("recommended_resources", [])
            if resources:
                _pdf_bold_line(pdf, "Resources:")
                for r in resources:
                    _pdf_bullet(pdf, r)

            pdf.ln(2)

    # 整体时间线
    timeline = report_data.get("overall_timeline", "")
    if timeline:
        _pdf_section_header(pdf, "Overall Timeline")
        _pdf_body(pdf, timeline)

    # 优先关注点
    focus = report_data.get("priority_focus", "")
    if focus:
        _pdf_section_header(pdf, "Priority Focus")
        _pdf_body(pdf, focus)

    # 顾问建议
    notes = report_data.get("advisor_notes", "")
    if notes:
        _pdf_section_header(pdf, "Advisor Notes")
        _pdf_body(pdf, notes)

    return _pdf_bytes(pdf)


def _pdf_bytes(pdf) -> bytes:
    """FPDF → bytes"""
    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf.read()


def _build_unified_pdf(report, report_data, lead):
    """生成统一学业风险与规划报告 PDF（使用 fpdf）"""
    pdf = FPDF()
    pdf.add_page()
    cn_font_path = _find_chinese_font()
    if cn_font_path:
        pdf.add_font("cn", "", cn_font_path, uni=True)
        pdf._cn_font = "cn"
    else:
        pdf._cn_font = "Helvetica"
    pdf.set_auto_page_break(auto=True, margin=20)

    def w(s):
        """写一行（自动折行）"""
        pdf.multi_cell(0, 8, s, new_x="LMARGIN", new_y="NEXT")

    def heading(s):
        pdf.set_font("cn", "", 14)
        pdf.set_text_color(30, 30, 80)
        pdf.cell(0, 12, s, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    def body(s):
        pdf.set_font("cn", "", 10)
        pdf.set_text_color(60, 60, 60)
        w(s)

    if not report_data:
        body("报告数据为空")
        return _pdf_bytes(pdf)

    pdf.set_font("cn", "", 18)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 15, "学业风险与规划报告", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("cn", "", 9)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, f"报告日期：{str(report.get('created_at', ''))[:10]}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    # 学生基础画像
    sp = report_data.get("student_profile", {})
    if sp:
        heading("学生基础画像")
        body(sp.get("background", ""))
        body(f"目标：{sp.get('target', '')}")
        ki = sp.get("key_info", {})
        for k, v in ki.items():
            body(f"{k}：{v}")
        pdf.ln(3)

    # 目标院校概览
    po = report_data.get("program_overview", "")
    if po:
        heading("目标院校专业概览")
        body(po)
        pdf.ln(3)

    # 逐课程分析
    courses = report_data.get("course_analysis", [])
    if courses:
        heading("核心课程逐门分析")
        for c in courses:
            # 检查是否需要新页
            if pdf.get_y() > pdf.h - 40:
                pdf.add_page()
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(30, 30, 80)
            name = c.get("course_name", "")
            credits = c.get("credits", "")
            pdf.cell(0, 7, f"{name}" + (f" ({credits})" if credits else ""), new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(60, 60, 60)
            assessment = c.get("assessment", "")
            if assessment:
                body(f"考核形式：{assessment}")
            prereq = c.get("prerequisites", "")
            if prereq:
                body(f"前置要求：{prereq}")
            readiness = c.get("student_readiness", "")
            if readiness:
                body(f"学生掌握：{readiness}")
            gap = c.get("gap", "")
            if gap:
                body(f"差距分析：{gap}")
            action = c.get("supplement_action", "")
            if action:
                pdf.set_text_color(80, 60, 160)
                body(f"建议补充：{action}")
                pdf.set_text_color(60, 60, 60)
            pdf.ln(2)

    # 差距分析
    gaps = report_data.get("gap_analysis", [])
    if gaps:
        heading("多维差距分析")
        for g in gaps:
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(30, 30, 80)
            pdf.cell(0, 8, f"{g.get('dimension', '')}（{g.get('risk', 'medium')}）", new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(60, 60, 60)
            body(f"当前：{g.get('current', '')}")
            body(f"要求：{g.get('required', '')}")
            body(f"差距：{g.get('gap', '')}")
            if g.get("improvement_strategy"):
                body(f"改进策略：{g['improvement_strategy']}")
            pdf.ln(2)

    # 准备计划
    plans = report_data.get("preparation_plan", [])
    if plans:
        heading("准备计划")
        for plan in plans:
            tl = plan.get("timeline", "")
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(30, 30, 80)
            label = plan.get("phase", "") + (f" · {tl}" if tl else "")
            pdf.cell(0, 8, label, new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(60, 60, 60)
            if plan.get("goal"):
                body(f"目标：{plan['goal']}")
            for task in plan.get("tasks", []):
                body(f"  - {task}")
            pdf.ln(2)

    # 综合评估
    oa = report_data.get("overall_assessment", "")
    if oa:
        heading("综合评估")
        body(oa)
        pdf.set_font("cn", "", 10)
        pdf.set_text_color(30, 30, 80)
        pdf.cell(0, 8, f"整体风险：{report_data.get('risk_level', 'medium')}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    # 建议
    recs = report_data.get("recommendations", [])
    if recs:
        heading("具体建议")
        for i, rec in enumerate(recs, 1):
            text = rec if isinstance(rec, str) else rec.get("action", str(rec))
            body(f"{i}. {text}")

    return _pdf_bytes(pdf)



def _build_growth_pdf(report, report_data, lead):
    """生成整体学情报告 PDF"""
    pdf = FPDF()
    pdf.add_page()
    cn_font_path = _find_chinese_font()
    if cn_font_path:
        pdf.add_font("cn", "", cn_font_path, uni=True)
        pdf._cn_font = "cn"
    else:
        pdf._cn_font = "Helvetica"
    pdf.set_auto_page_break(auto=True, margin=20)
    
    def heading(s):
        pdf.set_font("cn", "", 16)
        pdf.set_text_color(30, 30, 80)
        pdf.cell(0, 12, s, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)
    
    def subheading(s):
        pdf.set_font("cn", "", 12)
        pdf.set_text_color(30, 64, 175)
        pdf.cell(0, 10, s, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    
    def body(s):
        pdf.set_font("cn", "", 10)
        pdf.set_text_color(60, 60, 60)
        pdf.multi_cell(0, 7, s)
        pdf.ln(2)
    
    def bold_line(label, value):
        pdf.set_font("cn", "", 10)
        pdf.set_text_color(30, 30, 80)
        pdf.cell(0, 7, f"{label}：{value}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
    
    pdf.set_font("cn", "", 20)
    pdf.set_text_color(30, 30, 80)
    pdf.cell(0, 15, report_data.get("report_title", "阶段性学情综合报告"), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)
    
    si = report_data.get("student_info", {})
    if si:
        heading("学生信息")
        for k, v in si.items():
            bold_line(k, str(v))
    
    bl = report_data.get("baseline", "")
    if bl:
        heading("入学前基础概况")
        body(bl)
    
    ls = report_data.get("learning_summary", "")
    if ls:
        heading("学习历程总结")
        body(ls)
    
    for title, field in [("进步方面", "progress"), ("待改进方面", "weaknesses")]:
        items = report_data.get(field, [])
        if items:
            heading(title)
            for item in items:
                pdf.set_font("cn", "", 10)
                pdf.set_text_color(30, 30, 80)
                pdf.cell(0, 7, f"{item.get('aspect','')}：{item.get('detail','')}", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
    
    et = report_data.get("exam_trend", "")
    if et:
        heading("考试成绩趋势")
        body(et)
    
    dims = report_data.get("dimensions", {})
    if dims:
        heading("多维度评估")
        labels = {"academic":"学术","language":"语言","psychology":"心理","adaptation":"适应","planning":"规划"}
        for key, dim in dims.items():
            score = dim.get("score", "")
            note = dim.get("note", "")
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(30, 30, 80)
            pdf.cell(0, 7, f"{labels.get(key,key)}（{score}）", new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("cn", "", 9)
            pdf.set_text_color(60, 60, 60)
            if note:
                pdf.multi_cell(0, 6, note)
            pdf.ln(1)
    
    oa = report_data.get("overall_assessment", "")
    if oa:
        heading("综合评估")
        body(oa)
    
    recs = report_data.get("recommendations", [])
    if recs:
        heading("后续建议")
        for i, rec in enumerate(recs, 1):
            pdf.set_font("cn", "", 10)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(0, 7, f"{i}. {rec}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(1)
    
    return _pdf_bytes(pdf)

def generate_pdf(report, lead, report_type="risk") -> bytes:
    """
    生成 PDF 文档

    Args:
        report: consulting_reports 行 dict（含 report_json）
        lead: leads 行 dict
        report_type: 'risk' | 'preparation' | 'unified'

    Returns:
        PDF 文件 bytes
    """
    report_data = _parse_report_data(report.get("report_json", ""))

    if report_type == "preparation":
        return _build_preparation_pdf(report, report_data, lead)
    if report_type == "growth_overall":
        return _build_growth_pdf(report, report_data, lead)
    if report_type == "unified":
        return _build_unified_pdf(report, report_data, lead)
    return _build_risk_pdf(report, report_data, lead)
